from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_images_with_captions(doc, images_with_captions):
    # Cria uma tabela com 2 colunas e o número de linhas necessário
    rows = (len(images_with_captions) + 1) // 2  # arredonda para cima
    table = doc.add_table(rows=rows, cols=2)

    for index, (image_path, caption_text) in enumerate(images_with_captions):
        cell = table.cell(index // 2, index % 2)  # Determine a célula correta
        # Adiciona a imagem
        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(5))
        
        # Alinha a imagem ao centro
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Adiciona a descrição
        caption_paragraph = cell.add_paragraph(caption_text)
        caption_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Cria um novo documento
doc = Document()

# Lista de imagens e suas respectivas descrições
images_with_captions = [
    (r"caminho_da_imagem", "Descrição da imagem 1"),
    (r"caminho_da_imagem", "Descrição da imagem 2"),
    (r"caminho_da_imagem", "Descrição da imagem 3"),
    (r"caminho_da_imagem", "Descrição da imagem 4"),
]

# Adiciona as imagens e as descrições
add_images_with_captions(doc, images_with_captions)

# Salva o documento
doc.save("documento_com_imagens.docx")
