from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


# Criar um novo documento
document = Document()

# Acessar a primeira seção do documento
section = document.sections[0]

# Acessar o cabeçalho da seção
header = section.header
header_paragraph = header.paragraphs[0]

# Adicionar uma linha embaixo do cabeçalho (bordas do parágrafo)
p = header_paragraph._element
pPr = p.get_or_add_pPr()
borders = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')   # Tipo de linha (simples)
bottom.set(qn('w:sz'), '12')        # Espessura da linha
bottom.set(qn('w:space'), '1')      # Espaçamento entre texto e linha
bottom.set(qn('w:color'), '000000') # Cor da linha (preto)
borders.append(bottom)
pPr.append(borders)

# Função auxiliar para criar o campo no rodapé
def create_field(run, instr_text):
    # Adicionar o campo de início
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(ns.qn('w:fldCharType'), 'begin')  # Usar o namespace corretamente
    run._r.append(fld_char_begin)

    # Adicionar a instrução do campo
    instr = OxmlElement('w:instrText')
    instr.text = instr_text  # Instrução de campo
    run._r.append(instr)

    # Adicionar o campo de fim
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(ns.qn('w:fldCharType'), 'end')  # Usar o namespace corretamente
    run._r.append(fld_char_end)

# Acessar o rodapé da primeira seção
section = document.sections[0]
footer = section.footer
footer_paragraph = footer.paragraphs[0]

# Adicionar o campo de número da página
run = footer_paragraph.add_run("Página ")
create_field(run, 'PAGE')

# Adicionar o texto " de " entre os números de página
footer_paragraph.add_run(" de ")

# Adicionar o campo de total de páginas
run = footer_paragraph.add_run()
create_field(run, 'NUMPAGES')

for _ in range(5):
    document.add_paragraph()


# Adicionar um parágrafo ao cabeçalho
header_paragraph = header.paragraphs[0]
header_paragraph.text = "SERGIO SOUZA DOS SANTOS JUNIOR\nENGENHEIRO MECANICO\nCREA 101810-5D/SC"

# Adicionar um título
titulo = document.add_heading('LAUDO TÉCNICO\nGUINDASTE ARTICULADO HIDRÁULICO\nLAUDO N°051/2023\nART Nº AC20230087880\n')
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Acessar o primeiro "run" do parágrafo do título
run = titulo.runs[0]

# Definir o tamanho da fonte
run.font.size = Pt(24)  # Aumenta o tamanho da fonte para 24 pontos

for _ in range(5):
    document.add_paragraph()

# Adicionar um parágrafo
document.add_paragraph('Proprietário: HORIZONTE DA AMAZONIA LOGISTICA LTDA')
document.add_paragraph('Execução: ENG.MECÂNICO SERGIO SOUZA DOS SANTOS JUNIOR')


# Inserir uma quebra de página antes do próximo conteúdo
document.add_page_break()

# Adicionar um parágrafo vazio (linha em branco)
document.add_paragraph()


document.add_heading('1. OBJETIVO', level=1)
paragraph_objetivo = document.add_paragraph("           Apresentar a empresa contratante(HORIZONTE DA AMAZONIA LOGISTICA LTDA - CNPJ 07.462.961/0004-00) Laudo Técnico que comprove as condições do Veículo e Equipamento de Guindar.")
paragraph_objetivo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 

document.add_heading('2. METODOLOGIA', level=1)
paragraph_metodologia01 = document.add_paragraph("            O presente laudo tem caráter formalizar as responsabilidades técnicas na execução das atividades, bem como visa a apresentação de procedimentos de segurança que atendem os requisitos mínimos de segurança fixados em norma, em especial o Anexo XII na Norma Regulamentadora N.º 12 - Segurança no Trabalho em Máquinas e Equipamentos e NBR 14.768 -  Guindaste articulado hidráulico Requisitos")
paragraph_metodologia01.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

document.add_paragraph("            O referido Laudo terá como planejamento as seguintes etapas: ")

# Adicionar uma lista
document.add_paragraph('Vistoria de segurança do Equipamento de Guindar; ', style='List Bullet')
document.add_paragraph('Registro fotográfico do Veículo e do Equipamento de Guindar; ', style='List Bullet')
document.add_paragraph('Ensaios não Destrutivos (END); ', style='List Bullet')


paragraph_metodologia02 = document.add_paragraph("            Este Laudo e os trabalhos estarão sob responsabilidade técnica do Engº de Mecânico Sérgio Souza dos Santos Junior, CREA 101810-5D/SC.  ")
paragraph_metodologia02.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

paragraph_metodologia03 = document.add_paragraph("            O Laudo terá/contará com a presença física de profissional capacitado em movimentação de carga desde o planejamento até a conclusão. ")
paragraph_metodologia03.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

document.add_page_break()
document.add_paragraph()

document.add_heading('3. IDENTIFICAÇÃO DA EMPRESA', level=1)

document.add_heading('2. VISTORIA DE SEGURANÇA NO EQUIPAMENTO DE GUINAR', level=1)

paragraph_VISTORIA01 = document.add_paragraph('            Visando apresentar todo o parâmetro de segurança adota pela empresa HORIZONTE DA AMAZONIA LOGISTICA LTDA - CNPJ 07.462.961/0004-00, apresentaremos a segue Check list e vistoria no veículo PLACA PMD7520 e no equipamento de guindar MUNCK ARGOS AGI 16.5-13.6/32.')
paragraph_VISTORIA01.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY




# Adicionar uma tabela com 10 linhas e 2 colunas
table = document.add_table(rows=12, cols=2)

# Mesclar todas as células da primeira linha em uma única célula
first_row = table.rows[0]
first_row.cells[0].merge(first_row.cells[1])

# Adicionar um parágrafo na célula mesclada com o estilo Heading 2
paragraph = first_row.cells[0].add_paragraph('IDENTIFICAÇÃO DO EQUIPAMENTO', style='Heading 2')

# Aplicar negrito e tamanho da fonte ao título da primeira linha
run = paragraph.runs[0]
run.bold = True
run.font.size = Pt(12)  # Ajustar o tamanho da fonte (opcional)

# Centralizar o título da primeira linha
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Preencher as células da tabela com conteúdo diferente em cada célula
conteudo = [
    ['Placa de sinalização ', 'Segue abaixo alguma placas de sinalização e instruções de segurança '],
    ['Veículo de transporte ', 'VW/13.190 WORKER - PLACA PMD7520 '],
    ['Nome do fabricante e marca: ARGOS ', 'Modelo: MUNCK ARGOS AGI 16.5-13.6/32 '],
    ['Data de fabricação (mês e ano); ', '--/2014 '],
    ['Número de série; ', '066002'],
    ['Modelo e/ou tipo; ', 'AGI 16.5-13.6/32 '],
    ['Alcance máximo vertical, em metros (mm); ', 'Guindaste hidráulico articulado modelo AGI 16.513.6/32, como momento de carga útil de 16.500kgm, com três lanças hidráulicas e duas lança manual, alcance horizontal hidráulico de 13.600 mm, alcance máximo horizontal 15.000 mm, alcance máximo vertical 18.200 mm.'],
    ['Capacidade nominal de carga, em quilonewtons-metro ou toneladas-metro (kN.m ou t.m); ', 'Momento de carga útil de 16.500 kgfm'],
    ['Tabela de capacidade de carga, em quilogramas ou toneladas (kg ou t), e alcance horizontal, em metros (m); ', 'Vide item anterior'],
    ['Altura máxima para transporte (posição de recolhimento da lança), em metros (m); ', 'Alcance horizontal hidráulico de 13.600mm '],
    ['Máxima rotação da coluna em graus;', '368º em giro.'],
]

# Iterar sobre as linhas e colunas (exceto a primeira linha)
for i in range(1, 12):  # Linhas de 1 a 11
    row = table.rows[i]
    for j in range(2):  # Colunas de 0 a 1
        cell = row.cells[j]
        paragraph = cell.paragraphs[0]
        paragraph.text = conteudo[i-1][j]  # Preencher com conteúdo específico
        
        # Justificar o conteúdo do parágrafo
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Adicionar bordas à tabela
def set_borders(table):
    tbl = table._tbl  # Acessa o elemento XML subjacente da tabela
    for cell in tbl.iter_tcs():  # Itera sobre todas as células
        tcPr = cell.tcPr  # Acessa as propriedades da célula
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            cell.insert(0, tcPr)
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  # Define como linha única
            border.set(qn('w:sz'), '4')  # Define a espessura da borda
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Cor preta
            tcBorders.append(border)
        tcPr.append(tcBorders)

set_borders(table)  # Aplica bordas à tabela


# Adicionar uma tabela com 10 linhas e 2 colunas
table = document.add_table(rows=27, cols=3)

# Mesclar todas as células da primeira linha em uma única célula
first_row = table.rows[0]
first_row.cells[0].merge(first_row.cells[1]).merge(first_row.cells[2])

# Adicionar um parágrafo na célula mesclada com o estilo Heading 2
paragraph = first_row.cells[0].add_paragraph('CHECK-LIST DE VERIFICAÇÃO', style='Heading 2')

# Aplicar negrito e tamanho da fonte ao título da primeira linha
run = paragraph.runs[0]
run.bold = True
run.font.size = Pt(12)  # Ajustar o tamanho da fonte (opcional)

# Centralizar o título da primeira linha
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Preencher as células da tabela com conteúdo diferente em cada célula
conteudo = [
    ['A base foi fixada ao chassi ou sobre chassi através de dispositivo próprio com resistência mecânica compatível com os esforços exigidos pelo conjunto, quando em operação, dentro dos raios e capacidades de cargas especificadas pelos fabricantes.', 'APROVADO', ''],
    ['O braço foram construído em aço estrutural de alta resistência mecânica, proporcionando ao conjunto a possibilidade de efetuar movimentos de elevação através do acionamento de cilindros hidráulicos. ', 'APROVADO', ''],
    ['A lança foi construída em aço estrutural de alta resistência mecânica, compatível com os esforços envolvidos na operação do equipamento de acordo com as especificações do fabricante, dotada de extensões telescópicas com placas de deslizamentos ou rolamentos e acionamento hidráulico.', 'APROVADO', ''],
    ['O sistema de giro foi construído em aço de alta resistência mecânica, compatível com as especificações do equipamento, possuir sistema de lubrificação e ser acionado hidraulicamente.', 'APROVADO', ''],
    ['Conjunto de sapatas extensíveis horizontalmente, incorporado à base do guindaste, com acionamento manual ou hidráulico e cilindros hidráulicos para patolamento na vertical, equipados com válvulas de segurança contra os efeitos da ruptura de mangueiras ou tubos.', 'APROVADO', ''],
    ['O sistema de estabilização garante total estabilidade do conjunto quando em operação, permitindo ao operador executar os trabalhos com total segurança, respeitados os limites especificados pelos gráficos de carga do equipamento?', 'APROVADO', ''],
    ['Possui sapata estabilizadora traseira adicional, de acordo com especificações de cada fabricante, quando a sapata estabilizadora incorporada ao guindaste não for suficiente para garantir a estabilidade do conjunto. ', 'APROVADO', ''],
    ['O gancho foi fabricado em aço de alta resistência, compatível com as cargas especificadas para cada equipamento, devendo possuir trava de segurança (para cabo de aço, cintas, etc.). ', 'APROVADO', ''],
    ['Os componentes hidráulicos possui resistência mínima de ruptura igual a quatro vezes a pressão de trabalho, bem como aqueles normalmente especificados em função do limite de ruptura, tais como mangueiras, tubos e conexões, que devem ter resistência de ruptura três vezes o valor da pressão de trabalho.', 'APROVADO', ''],
    ['Possui plaqueta/tabela de identificação das capacidade nominal de carga e momentos em cada estágio dos braços, das lanças e das extensões telescópicas.', 'APROVADO', ''],
    ['A alimentação do sistema hidráulico deve ser feita pela bomba hidráulica, acionada através da tomada de força compatível com o veículo, acoplada à caixa de mudança do veículo, diretamente ou através de eixo cardan, ou ainda através de motor auxiliar independente ou através de embreagem eletromagnética. ', 'APROVADO', ''],
    ['O acionamento da tomada de força deve ser efetuado através de um dispositivo de comando instalado na cabina do veículo, com indicação visual que sinalize o funcionamento. ', 'APROVADO', ''],
    ['O comando para todas as funções do equipamento devidamente identificado (elevação, lança, giro, estabilização, etc.) é efetuado através de distribuidor hidráulico dotado de válvulas reguladoras de pressão de carga e posicionado no equipamento de forma a permitir que o operador possa acionálo com total segurança de ambos os ladosdo veículo. A válvula de comando pode ser acionada manualmente e/ou através de controle remoto a cabo ou a rádio.', 'APROVADO', ''],
    ['Os cilindros hidráulicos devem são de duplo efeito, dotados de válvulas de segurança, fixados nos pontos de articulação através de pinos e buchas de resistência mecânica compatível com os esforços envolvidos na operação do guindaste articulado hidráulico, e dotados de sistema de lubrificação. ', 'APROVADO', ''],
    ['O guindaste articulado hidráulico possui válvulas de segurança em todo o sistema hidráulico, protegendo-o contra sobrepressões (válvula de alívio), sobrecargas, rupturas de mangueiras ', 'APROVADO', ''],
    ['Possui válvulas de segurança nos cilindros das sapatas estabilizadoras contra os efeitos da ruptura de tubos e mangueiras; ', 'APROVADO', 'Todos os pontos hidráulicos possuem válvula de segurança contra perda de pressão e contra-balanço. Em caso de sobre carga está estabiliza e se mantém imóvel até retenção da carga especificada pelo fabricante. '],
    ['Possui válvulas de pressão de carga, nos cilindros de elevação (coluna e braço) e do braço (braço e lança), com a função de evitar operações fora das especificadas no gráfico de cargas de cada guindaste articulado hidráulico, de acordo com as especificações do fabricante; ', 'APROVADO', ''],
    ['Possui válvulas de segurança instaladas nos cilindros da lança telescópica e extensões hidráulicas, com a função de bloqueio instantâneo das operações em caso de rupturas de mangueiras e tubulações ', 'APROVADO', ''],
    ['Possui válvulas de alívio e reguladoras de pressão de carga, instaladas no comando hidráulico', 'APROVADO', ''],
    ['O reservatório de óleo possui indicador de nível de óleo mínimo e máximo?', 'APROVADO', ''],
    ['Possui respiro, devidamente protegido contra a entrada de água e poeira;', 'APROVADO', ''],
    ['Possui bocal de enchimento com tela de proteção; ', 'APROVADO', ''],
    ['Possui sistema de identificação do óleo utilizado? ', 'APROVADO', ''],
    ['O chassi auxiliar foi acoplado às longarinas do veículo em toda a sua extensão através de sistema de fixação convencional destinado a tal fim, conforme especificações dos fabricantes e/ou das montadoras de veículos. O chassi auxiliar deve ser dimensionado de tal forma que, somado ao chassi do veículo, venha a possuir resistência mecânica compatível para o uso do guindaste em sua plenitude.', 'APROVADO', ''],
    ['O Guindaste permite recolhimento do braço, da lança e das extensões telescópicas que, através de articulações, se acomodem totalmente no dispositivo de apoio, sem interferir com a geometria do veículo?', 'APROVADO', ''],
    ['A base e a coluna foram construídas em aço estrutural de alta resistência mecânica, compatíveis com as especificações e características do respectivo equipamento?', 'APROVADO', ''],
]

# Iterar sobre as linhas e colunas (exceto a primeira linha)
for i in range(1, 26):  # Linhas de 1 a 11
    row = table.rows[i]
    for j in range(3):  # Colunas de 0 a 2
        cell = row.cells[j]
        paragraph = cell.paragraphs[0]
        paragraph.text = conteudo[i-1][j]  # Preencher com conteúdo específico
        
        # Verificar se a célula está na coluna do meio (índice 1)
        if j == 1:
            # Centralizar o conteúdo da coluna do meio
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Alterar o fundo da célula para verde
            shading_elm = parse_xml(r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls('w')))
            cell._element.get_or_add_tcPr().append(shading_elm)

        else:
            # Justificar o conteúdo das colunas laterais
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

set_borders(table)  # Aplica bordas à tabela

# Salvar o documento
document.save('meu_documento.docx')
